from __future__ import annotations

from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def export_docx(docs_dir: Path, output: Path) -> None:
    """Export markdown files from docs_dir to a single Word document."""
    if not HAS_DOCX:
        raise ImportError(
            "python-docx is required for Word export. "
            "Install with: pip install codemap[docx]"
        )

    doc = Document()

    md_files = sorted(docs_dir.glob("*.md"))
    if not md_files:
        doc.add_paragraph("No documents found.")
        doc.save(str(output))
        return

    for i, md_file in enumerate(md_files):
        if i > 0:
            doc.add_page_break()

        md_text = md_file.read_text(encoding="utf-8")
        _parse_markdown_to_docx(doc, md_text)

    doc.save(str(output))


def _parse_markdown_to_docx(doc: Document, md_text: str) -> None:
    """Simple markdown to docx conversion."""
    for line in md_text.strip().split("\n"):
        stripped = line.strip()
        if not stripped:
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped.startswith("|"):
            # Skip table separator lines and handle table rows
            if stripped.startswith("|--") or stripped.startswith("| --"):
                continue
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            if cells:
                doc.add_paragraph(" | ".join(cells))
        else:
            doc.add_paragraph(stripped)
